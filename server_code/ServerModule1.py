import anvil.users
import anvil.tables as tables
import anvil.tables.query as q
from anvil.tables import app_tables
import anvil.server
import anvil.media
import torchxrayvision as xrv
import skimage, torch, torchvision
from docx import Document
from docx.shared import Inches
import datetime

# To allow anvil.server.call() to call functions here, we mark
# them with @anvil.server.callable.
# Here is an example - you can replace it with your own:
#
# @anvil.server.callable
# def say_hello(name):
#   print("Hello, " + name + "!")
#   return 42
#
@anvil.server.callable
def get_models():
  return app_tables.models.search()

@anvil.server.callable
def get_model(name):
  return app_tables.models.get(name=name)

@anvil.server.callable
def classify_image(file, modelname):
  with anvil.media.TempFile(file) as filename:
    #img = load_img(filename)
    img = skimage.io.imread(filename)
    img = xrv.datasets.normalize(img, 255) # convert 8-bit image to [-1024, 1024] range
    img = img.mean(2)[None, ...] # Make single color channel
    transform = torchvision.transforms.Compose([xrv.datasets.XRayCenterCrop(),xrv.datasets.XRayResizer(224)])
    img = transform(img)
    img = torch.from_numpy(img)
    model = xrv.models.DenseNet(weights=modelname)
    outputs = model(img[None,...]) # or model.features(img[None,...]) 
    return str(dict(zip(model.pathologies,outputs[0].detach().numpy())))

@anvil.server.callable
def getHeatmap(file, modelname, patologie):
  with anvil.media.TempFile(file) as filename:
    #img = load_img(filename)
    img = skimage.io.imread(filename)
    img = xrv.datasets.normalize(img, 255) # convert 8-bit image to [-1024, 1024] range
    #img = img.mean(2)[None, ...] # Make single color channel
    if len(img.shape) > 2:
      img = img[:, :, 0]
    if len(img.shape) < 2:
      print("error, dimension lower than 2 for image")
    img = img[None, :, :]
    transform = torchvision.transforms.Compose([xrv.datasets.XRayCenterCrop(),xrv.datasets.XRayResizer(224)])
    img = transform(img)
    img = torch.from_numpy(img)
    model = xrv.models.DenseNet(weights=modelname)
    target = model.pathologies.index(patologie)
    img = img.requires_grad_()
    outputs = model(img[None,...]) # or model.features(img[None,...]) 
    grads = torch.autograd.grad(outputs[:,target], img)[0][0][0]
    blurred = skimage.filters.gaussian(grads.detach().cpu().numpy()**2, sigma=(5, 5), truncate=3.5)

    my_dpi = 100
    fig = plt.figure(frameon=False, figsize=(224/my_dpi, 224/my_dpi), dpi=my_dpi)
    ax = plt.Axes(fig, [0., 0., 1., 1.])
    ax.set_axis_off()
    fig.add_axes(ax)
    ax.imshow(img[0][0].detach().cpu().numpy(), cmap="gray", aspect='auto')
    ax.imshow(blurred, alpha=0.5)
    ax.savefig('temp.png')
    return True

@anvil.server.callable
def save_report(file, utilizator, modelname, result):
    print(f'Utilizator: {utilizator}; Model: {modelname}')
    doc = Document()
    doc.add_heading('Raport analiză {file.name}', 0)
    doc.add_paragraph('Configurație:')
    doc.add_paragraph('Utilizator:{utilizator}', style='List Number')
    doc.add_paragraph('Data:{datetime.now()}', style='List Number')
    doc.add_paragraph('Model:{modelname}', style='List Number')
    with anvil.media.TempFile(file) as filename:
      doc.add_picture(filename, width=Inches(3))
    doc.add_paragraph('Figura 1: Imaginea originală')
    keysList = list(eval(result).keys())
    item_list = [{'vname':name, 'vvalue':eval(result)[name]} for name in keysList]
    for item in item_list:
      doc.add_paragraph(f"{item['vname']}: {item['vvalue']}', style='List Number'")
      if item['vvalue'] > 0.5:
        print(f"HM Model:{item['vname']} ({item['vvalue']})")
        if getHeatmap(file=file, modelname=modelname, patologie=item['vname']):
          doc.add_picture('temp.png', width=Inches(3))
    doc.save(f'{utilizator}_{modelname}_{file.name}.docx')
  